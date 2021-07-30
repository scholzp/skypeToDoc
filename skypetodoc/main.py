import tarfile
import json
import docx
import argparse
import re
from docx import Document



def extractChats(fn):
    tarObj = tarfile.open(fn)
    tarObj.extractall("./tmp/chats")

    return "./tmp/chats/messages.json"


def loadJson(fn):
    file = open(fn)
    jsonContent = json.load(file)
    file.close()
    return jsonContent


def printExportInfo(content):
    print('User ID:', content['userId'])
    print('Date of Export:', content['exportDate'])


def extractChatsFromJson(json_string):
    conversationList = json_string['conversations']
    resultList = []
    for chat in conversationList:
        if len(chat['MessageList']) != 0:
            resultList.append(chat)
    print('Total number of chats:', str(len(conversationList)))
    print('Number of empty chats:', str(len(conversationList) - len(resultList)))
    print('Numbe of exportable chats:', str(len(resultList)))
    return resultList


def printChatInfos(chatList):
    for index in range(1, len(chatList)):
        chat = chatList[index-1]
        threadProperties = chat['threadProperties']
        members = 'Not a group'
        topic = 'No topic available'

        if chat['threadProperties']:
            members = chat['threadProperties']['members']
            if chat['threadProperties']['topic']:
                topic = chat['threadProperties']['topic']

        print('{0}: Name:{1} Members:{2} topic:{3}'.format(str(index), chat['displayName'], members, topic))


def exportChat(chatList, index, dest):
    chat = chatList[index]
    exportAsDocx(chat, dest)


def exportAsDocx(chat, dest):
    document = docx.Document()
    for message in list(chat['MessageList'])[::-1]:
        author = message['from']
        content = message['content']
        content = re.sub(r'(<a href=.*\">)|(</a>)', '', content)
        content = re.sub(r'(<ss type=.*\">)|(</ss>)', '', content)
        if re.search(r'</legacyquote>.*<legacyquote>', content):
            quoteContent = re.search(r'</legacyquote>.*<legacyquote>', content).__getitem__(0)
            quoteContent = re.sub(r'(</legacyquote>)|(<legacyquote>)', '', quoteContent)
            quoteAuthor = re.sub(r'(author=)|(authorname)', '', re.search(r'author=\".*\" authorname', content).__getitem__(0))
            content = re.sub(r'<quote.*</legacyquote>', ( 'Quote ' + quoteAuthor + ': \n'), content, flags=re.DOTALL)
            content = re.sub(r'<legacyquote>.*</quote>', 'End of Qute', content, flags=re.DOTALL)
        time = message['originalarrivaltime']
        document.add_heading((author, ' wrote at ', time, ' : '), level=4)
        document.add_paragraph(content)
    document.save(dest)


def main():
    parser = argparse.ArgumentParser(
        description='Create documents from skype chats.')
    parser.add_argument('-input', required=True, action='store',
                        help='Input file. Must be path to the exported chats.')

    print('Initiating...')
    args = parser.parse_args()
    extractChats(args.input)
    content = loadJson("./tmp/chats/messages.json")
    printExportInfo(content)
    chatList = extractChatsFromJson(content)

    printChatInfos(chatList)
    print('\n Choose an chat index to export.')
    index = int(input()) - 1
    exportChat(chatList, index, 'test.docx')


if __name__ == "__main__":
    main()
